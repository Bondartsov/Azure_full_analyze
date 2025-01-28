from docx import Document
from core.logging.logger import log
import os

def generate_report(project_name, repository_name, total_tokens, total_commits, analysis):
    """
    Генерирует отчёт о репозитории, включая количество токенов и коммитов.
    """
    log("Создание отчёта...")

    doc = Document()
    doc.add_heading(f'Отчёт по проекту: {project_name}', level=1)

    doc.add_paragraph(f'📂 Репозиторий: {repository_name}')
    doc.add_paragraph(f'📊 Общее количество токенов: {total_tokens}')
    doc.add_paragraph(f'🔢 Общее количество коммитов: {total_commits}')

    # Добавляем топ-авторов, если они есть
    if analysis.get("top_authors"):
        doc.add_heading("Топ-5 авторов:", level=2)
        for author, commit_count in analysis["top_authors"]:
            doc.add_paragraph(f"{author}: {commit_count} коммитов")

    # Определяем путь к файлу отчёта
    reports_dir = os.path.join(os.path.dirname(__file__), "../../reports")
    os.makedirs(reports_dir, exist_ok=True)

    report_path = os.path.join(reports_dir, f"{project_name}_{repository_name}_report.docx")
    doc.save(report_path)

    log(f"Отчёт успешно сохранён: {report_path}")
    return report_path
