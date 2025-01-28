from docx import Document

def generate_report(project_name, repository_name, analysis):
    """
    Генерация отчёта в формате Word.
    """
    document = Document()
    document.add_heading(f'Отчёт по проекту: {project_name}', level=1)
    document.add_heading(f'Репозиторий: {repository_name}', level=2)

    document.add_paragraph(f"Общее количество коммитов: {analysis['total_commits']}")

    document.add_heading('Топ-5 авторов:', level=3)
    for author, count in analysis['top_authors']:
        document.add_paragraph(f"{author}: {count} коммитов")

    # Сохранение документа
    report_path = f"{project_name}_{repository_name}_report.docx"
    document.save(report_path)
    print(f"Отчёт сохранён: {report_path}")